"""
Excel转Word自动盖章程序
功能：
1. 生成测试Excel和印章图片
2. 读取Excel表格写入Word
3. 每页添加电子印章
依赖库安装：
pip install openpyxl python-docx Pillow faker
"""

import os
import random
from pathlib import Path
from faker import Faker
from PIL import Image, ImageDraw, ImageFont
import openpyxl
from openpyxl.styles import Font
from docx import Document
from docx.shared import Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----------------------
# 第一部分：生成测试文件
# ----------------------

def generate_test_files():
    """生成测试Excel和印章图片"""
    # 创建测试目录
    Path("test_files").mkdir(exist_ok=True)
    
    # 生成Excel测试数据
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "员工信息"
    
    # 设置表头
    headers = ["工号", "姓名", "部门", "薪资", "入职日期"]
    ws.append(headers)
    
    # 生成20条测试数据
    fake = Faker("zh_CN")
    for i in range(1, 21):
        ws.append([
            f"EMP{1000+i}",
            fake.name(),
            random.choice(["技术部", "市场部", "财务部", "人事部"]),
            round(random.uniform(8000, 25000), 2),
            fake.date_between(start_date="-5y")
        ])
    
    # 设置样式
    bold_font = Font(bold=True)
    for cell in ws[1]:
        cell.font = bold_font
    
    wb.save("test_files/test_data.xlsx")

    # 生成电子印章图片
    img = Image.new("RGBA", (400, 400), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    
    # 绘制圆形印章
    draw.ellipse((20, 20, 380, 380), outline="red", width=8)
    
    # 添加文字
    font = ImageFont.truetype("simhei.ttf", 40)
    draw.text((100, 150), "电子签章", fill="red", font=font)
    draw.text((120, 250), "专用章", fill="red", font=font)
    
    img.save("test_files/stamp.png")

# ----------------------
# 第二部分：核心处理逻辑
# ----------------------

def excel_to_word_with_stamp():
    """将Excel表格转换到Word并添加电子印章"""
    doc = Document()
    
    # 设置页面布局
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    
    # 读取Excel数据
    wb = openpyxl.load_workbook("test_files/test_data.xlsx")
    ws = wb.active
    
    # 处理数据
    for row_idx, row in enumerate(ws.iter_rows(values_only=True), 1):
        # 每5行数据创建一个表格页
        if (row_idx-1) % 5 == 0:
            if row_idx > 1:  # 添加分页符（第一页除外）
                doc.add_page_break()
            
            # 创建新表格
            table = doc.add_table(rows=6, cols=5)  # 5列，标题行+5数据行
            table.style = "Table Grid"
            
            # 设置表格列宽
            for col in table.columns:
                col.width = Cm(4)
            
            # 添加表头
            for col_idx, header in enumerate(ws[1]):
                table.cell(0, col_idx).text = str(header)
            
            # 添加印章（使用绝对定位）
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture("test_files/stamp.png", width=Cm(3.5))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
        # 填充表格数据
        if row_idx > 1:  # 跳过标题行
            current_row = (row_idx-2) % 5 + 1
            for col_idx, value in enumerate(row):
                table.cell(current_row, col_idx).text = str(value)

    # 保存文档
    doc.save("test_files/output.docx")

# ----------------------
# 执行主程序
# ----------------------
if __name__ == "__main__":
    if not os.path.exists("test_files/test_data.xlsx"):
        generate_test_files()
        print("测试文件已生成在 test_files 目录")
    
    excel_to_word_with_stamp()
    print("处理完成，输出文件：test_files/output.docx")